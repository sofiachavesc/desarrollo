from docx import Document

# Ruta del archivo del orden del día (opcional)
ORDEN_DIA = [
    "Apertura de la asamblea",
    "Lectura y aprobación del acta anterior",
    "Discusión de temas generales",
    "Votación sobre propuestas",
    "Cierre de la asamblea"
]

# Crear el documento de Word
def crear_acta(transcripcion, orden_dia):
    try:
        # Crear un nuevo documento
        doc = Document()

        # Encabezado
        doc.add_heading("ASAMBLEA GENERAL ORDINARIA", level=1).alignment = 1  # Centrado
        doc.add_paragraph("NOMBRE DE LA ENTIDAD").bold = True
        doc.add_paragraph("ACTA No: ___________").bold = True
        doc.add_paragraph("ORDEN DEL DÍA:").bold = True

        # Agregar el orden del día
        for i, punto in enumerate(orden_dia, start=1):
            doc.add_paragraph(f"{i}. {punto}")

        doc.add_paragraph("Hora de inicio: ____________________\n")

        # Procesar transcripción
        doc.add_heading("DESARROLLO DE LA ASAMBLEA", level=2)

        # Dividir la transcripción en líneas
        lineas = transcripcion.split("\n")
        for linea in lineas:
            if "votar" in linea.lower() or "resultados" in linea.lower():
                # Resaltar votaciones
                doc.add_paragraph(linea).bold = True
            elif any(punto.lower() in linea.lower() for punto in orden_dia):
                # Resaltar puntos del orden del día
                doc.add_paragraph(linea).bold = True
            else:
                # Agregar texto normal
                doc.add_paragraph(linea)

        # Guardar el documento
        output_path = "C:\\Users\\sofia\\Desktop\\actas bot\\Acta_Asamblea.docx"
        doc.save(output_path)
        print(f"Acta generada correctamente: {output_path}")

    except Exception as e:
        print(f"Error al generar el acta: {e}")

# Llamada para generar el acta
if __name__ == "__main__":
    # Simula la transcripción como entrada
    transcripcion = """
    Apertura de la asamblea por Juan Pérez, Torre 1 Apartamento 101.
    Se procede a votar por la aprobación del acta anterior. Resultados: Aprobada por unanimidad.
    Discusión de temas generales sobre el presupuesto.
    Se abre la votación para aprobar la nueva propuesta. Resultados: 10 a favor, 2 en contra.
    Cierre de la asamblea.
    """
    crear_acta(transcripcion, ORDEN_DIA)
